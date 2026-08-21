"""
word_generator.py — Generates dispatch label Word files from extracted order data.

Creates .docx files with the following format per order (font size 14):
    Order Id: #DXXXXX  (bold, no indent)
    To                 (bold, no indent)
        Customer Name      (bold, indented)
        Address Line 1     (normal, indented)
        City, State - Pin  (normal, indented)
        Ph: XXXXXXXXXX     (Ph: bold, number normal, indented)

3 orders per page for bulk processing, with page breaks after every 3rd order.
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_dispatch_word(
    orders: list[dict],
    output_dir: str,
) -> str:
    """
    Generate a dispatch label Word file from a list of order dicts.

    Each order dict has:
        - order_id: str
        - customer_name: str
        - address_line1: str
        - address_line2: str (City, State - Pincode)
        - phone: str

    Args:
        orders: List of order data dicts
        output_dir: Directory to save the generated .docx file

    Returns:
        Full path to the generated .docx file
    """
    if not orders:
        raise ValueError("No order data provided")

    os.makedirs(output_dir, exist_ok=True)

    # Determine file name based on order IDs
    filename = _generate_filename(orders)
    output_path = os.path.join(output_dir, filename)

    # Create Word document
    doc = Document()

    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(14)

    # Set margins
    for section in doc.sections:
        section.top_margin = Pt(36)     # 0.5 inch
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)    # 0.75 inch
        section.right_margin = Pt(54)

    for i, order in enumerate(orders):
        is_first = (i % 3 == 0)
        _add_order_to_doc(doc, order, is_first_on_page=is_first)

        # Add spacing between orders on the same page (page break every 3rd order)
        if i < len(orders) - 1:
            if (i + 1) % 3 == 0:
                doc.add_page_break()

    doc.save(output_path)
    return output_path


def _add_order_to_doc(doc: Document, order: dict, is_first_on_page: bool = False):
    """Add a single dispatch label to the document."""

    # Indentation for customer details
    indent = Inches(0.5)

    # Line 1: "Order Id: #DXXXXX" (bold, no indent)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    # Add space before the order if it's not the first on the page
    para.paragraph_format.space_before = Pt(0) if is_first_on_page else Pt(36)
    
    run = para.add_run(f"Order Id: #D{order['order_id']}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    # Line 2: "To" (bold, no indent)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run("To")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    # Line 3: Customer Name (bold, indented)
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.left_indent = indent
    run = para.add_run(order["customer_name"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"

    # Line 4: Address Line 1 (normal, indented)
    if order.get("address_line1"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent = indent
        run = para.add_run(order["address_line1"])
        run.bold = False
        run.font.size = Pt(14)
        run.font.name = "Arial"

    # Line 5: City, State - Pincode (normal, indented)
    if order.get("address_line2"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent = indent
        
        # Format 2 requires a period at the end of the pincode line
        addr2 = order["address_line2"].strip()
        if not addr2.endswith('.'):
            addr2 += '.'
            
        run = para.add_run(addr2)
        run.bold = False
        run.font.size = Pt(14)
        run.font.name = "Arial"

    # Line 6: "Ph: XXXXXXXXXX" (Ph: bold, number normal, indented)
    if order.get("phone"):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.left_indent = indent
        # "Ph:" in bold
        run_label = para.add_run("Ph: ")
        run_label.bold = True
        run_label.font.size = Pt(14)
        run_label.font.name = "Arial"
        # Phone number in normal
        run_number = para.add_run(order["phone"])
        run_number.bold = False
        run_number.font.size = Pt(14)
        run_number.font.name = "Arial"


def _generate_filename(orders: list[dict]) -> str:
    """
    Generate output filename based on order IDs.

    Single order:  D15125.docx
    Bulk orders:   D15100_to_D15125.docx
    """
    if len(orders) == 1:
        return f"D{orders[0]['order_id']}.docx"
    else:
        # Sort by order_id numerically for proper range
        sorted_orders = sorted(orders, key=lambda o: int(o["order_id"]))
        first_id = sorted_orders[0]["order_id"]
        last_id = sorted_orders[-1]["order_id"]
        return f"D{first_id}_to_D{last_id}.docx"


if __name__ == "__main__":
    # Quick test
    test_orders = [
        {
            "order_id": "15125",
            "customer_name": "Sanvitha allam",
            "address_line1": "5-4-99/A , , Pakabanda bazar, khammam, Telegana",
            "address_line2": "Khammam, Telangana - 507001",
            "phone": "8074880903",
        },
        {
            "order_id": "15126",
            "customer_name": "Test User",
            "address_line1": "123 Main Street",
            "address_line2": "Hyderabad, Telangana - 500081",
            "phone": "9876543210",
        },
        {
            "order_id": "15127",
            "customer_name": "Another User",
            "address_line1": "456 Park Avenue",
            "address_line2": "Chennai, Tamil Nadu - 600001",
            "phone": "8765432109",
        },
        {
            "order_id": "15128",
            "customer_name": "Fourth User",
            "address_line1": "789 Lake Road",
            "address_line2": "Bangalore, Karnataka - 560001",
            "phone": "7654321098",
        },
    ]
    output = generate_dispatch_word(test_orders, "./test_output")
    print(f"Generated: {output}")
